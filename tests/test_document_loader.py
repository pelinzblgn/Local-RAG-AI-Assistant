from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from src.document_loader import (
    collect_document_files,
    collect_text_files,
    find_document_files,
    find_text_files,
    is_supported_document_file,
    is_supported_text_file,
    read_document_file,
    read_text_file,
)


def test_is_supported_text_file_accepts_txt() -> None:
    path = Path("notes.txt")

    assert is_supported_text_file(path) is True


def test_is_supported_text_file_rejects_other_extension() -> None:
    path = Path("notes.pdf")

    assert is_supported_text_file(path) is False


def test_is_supported_document_file_accepts_txt() -> None:
    assert is_supported_document_file(
        Path("notes.txt")
    ) is True


def test_is_supported_document_file_accepts_pdf() -> None:
    assert is_supported_document_file(
        Path("notes.pdf")
    ) is True


def test_is_supported_document_file_accepts_docx() -> None:
    assert is_supported_document_file(
        Path("notes.docx")
    ) is True


def test_is_supported_document_file_is_case_insensitive() -> None:
    assert is_supported_document_file(
        Path("NOTES.PDF")
    ) is True

    assert is_supported_document_file(
        Path("NOTES.DOCX")
    ) is True


def test_is_supported_document_file_rejects_unknown_extension() -> None:
    assert is_supported_document_file(
        Path("notes.xlsx")
    ) is False


def test_read_text_file_reads_utf8_content(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "notes.txt"
    file_path.write_text(
        "STM32 test content",
        encoding="utf-8",
    )

    content = read_text_file(file_path)

    assert content == "STM32 test content"


def test_read_text_file_rejects_missing_file(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "missing.txt"

    with pytest.raises(FileNotFoundError):
        read_text_file(file_path)


def test_read_text_file_rejects_unsupported_extension(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "notes.pdf"
    file_path.write_text(
        "Example",
        encoding="utf-8",
    )

    with pytest.raises(ValueError):
        read_text_file(file_path)


def test_read_text_file_rejects_empty_file(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "empty.txt"
    file_path.write_text(
        "",
        encoding="utf-8",
    )

    with pytest.raises(ValueError):
        read_text_file(file_path)


def test_read_document_file_reads_txt(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "document.txt"
    file_path.write_text(
        "LocalMind TXT document",
        encoding="utf-8",
    )

    content = read_document_file(file_path)

    assert content == "LocalMind TXT document"


def test_read_document_file_reads_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "document.docx"

    document = Document()
    document.add_paragraph(
        "LocalMind DOCX document"
    )
    document.add_paragraph(
        "Second paragraph"
    )
    document.save(file_path)

    content = read_document_file(file_path)

    assert "LocalMind DOCX document" in content
    assert "Second paragraph" in content


def test_read_document_file_rejects_empty_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "empty.docx"

    document = Document()
    document.save(file_path)

    with pytest.raises(ValueError):
        read_document_file(file_path)


def test_read_document_file_rejects_corrupt_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "broken.docx"
    file_path.write_bytes(
        b"not-a-valid-docx"
    )

    with pytest.raises(RuntimeError):
        read_document_file(file_path)


def test_read_document_file_rejects_pdf_without_text(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "empty.pdf"

    writer = PdfWriter()
    writer.add_blank_page(
        width=100,
        height=100,
    )

    with file_path.open("wb") as output:
        writer.write(output)

    with pytest.raises(ValueError):
        read_document_file(file_path)


def test_read_document_file_rejects_corrupt_pdf(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "broken.pdf"
    file_path.write_bytes(
        b"not-a-valid-pdf"
    )

    with pytest.raises(RuntimeError):
        read_document_file(file_path)


def test_read_document_file_rejects_unsupported_extension(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "notes.xlsx"
    file_path.write_text(
        "Unsupported",
        encoding="utf-8",
    )

    with pytest.raises(ValueError):
        read_document_file(file_path)


def test_find_text_files_finds_only_txt_files(
    tmp_path: Path,
) -> None:
    first_file = tmp_path / "a.txt"
    second_file = tmp_path / "b.txt"
    ignored_file = tmp_path / "c.pdf"

    first_file.write_text(
        "A",
        encoding="utf-8",
    )
    second_file.write_text(
        "B",
        encoding="utf-8",
    )
    ignored_file.write_text(
        "C",
        encoding="utf-8",
    )

    result = find_text_files(tmp_path)

    assert result == [
        first_file,
        second_file,
    ]


def test_find_document_files_finds_supported_documents(
    tmp_path: Path,
) -> None:
    txt_file = tmp_path / "a.txt"
    pdf_file = tmp_path / "b.pdf"
    docx_file = tmp_path / "c.docx"
    ignored_file = tmp_path / "d.xlsx"

    txt_file.write_text(
        "TXT",
        encoding="utf-8",
    )
    pdf_file.write_bytes(
        b"placeholder"
    )
    docx_file.write_bytes(
        b"placeholder"
    )
    ignored_file.write_text(
        "ignored",
        encoding="utf-8",
    )

    result = find_document_files(tmp_path)

    assert result == [
        txt_file,
        pdf_file,
        docx_file,
    ]


def test_find_text_files_is_non_recursive_by_default(
    tmp_path: Path,
) -> None:
    root_file = tmp_path / "root.txt"
    nested_directory = tmp_path / "nested"
    nested_directory.mkdir()

    nested_file = nested_directory / "nested.txt"

    root_file.write_text(
        "Root",
        encoding="utf-8",
    )
    nested_file.write_text(
        "Nested",
        encoding="utf-8",
    )

    result = find_text_files(tmp_path)

    assert result == [root_file]


def test_find_text_files_can_scan_recursively(
    tmp_path: Path,
) -> None:
    root_file = tmp_path / "root.txt"
    nested_directory = tmp_path / "nested"
    nested_directory.mkdir()

    nested_file = nested_directory / "nested.txt"

    root_file.write_text(
        "Root",
        encoding="utf-8",
    )
    nested_file.write_text(
        "Nested",
        encoding="utf-8",
    )

    result = find_text_files(
        tmp_path,
        recursive=True,
    )

    assert set(result) == {
        root_file,
        nested_file,
    }


def test_find_document_files_can_scan_recursively(
    tmp_path: Path,
) -> None:
    root_file = tmp_path / "root.txt"

    nested_directory = tmp_path / "nested"
    nested_directory.mkdir()

    nested_pdf = (
        nested_directory / "nested.pdf"
    )

    root_file.write_text(
        "Root",
        encoding="utf-8",
    )
    nested_pdf.write_bytes(
        b"placeholder"
    )

    result = find_document_files(
        tmp_path,
        recursive=True,
    )

    assert set(result) == {
        root_file,
        nested_pdf,
    }


def test_find_text_files_returns_empty_list_for_empty_directory(
    tmp_path: Path,
) -> None:
    result = find_text_files(tmp_path)

    assert result == []


def test_collect_text_files_accepts_single_txt_file(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "single.txt"
    file_path.write_text(
        "Single document",
        encoding="utf-8",
    )

    result = collect_text_files(file_path)

    assert result == [file_path]


def test_collect_document_files_accepts_single_pdf(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "single.pdf"
    file_path.write_bytes(
        b"placeholder"
    )

    result = collect_document_files(
        file_path
    )

    assert result == [file_path]


def test_collect_document_files_accepts_single_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "single.docx"
    file_path.write_bytes(
        b"placeholder"
    )

    result = collect_document_files(
        file_path
    )

    assert result == [file_path]


def test_collect_text_files_accepts_directory(
    tmp_path: Path,
) -> None:
    first_file = tmp_path / "first.txt"
    second_file = tmp_path / "second.txt"

    first_file.write_text(
        "First",
        encoding="utf-8",
    )
    second_file.write_text(
        "Second",
        encoding="utf-8",
    )

    result = collect_text_files(tmp_path)

    assert result == [
        first_file,
        second_file,
    ]


def test_collect_document_files_accepts_directory(
    tmp_path: Path,
) -> None:
    txt_file = tmp_path / "a.txt"
    pdf_file = tmp_path / "b.pdf"
    docx_file = tmp_path / "c.docx"
    ignored_file = tmp_path / "d.xlsx"

    txt_file.write_text(
        "TXT",
        encoding="utf-8",
    )
    pdf_file.write_bytes(
        b"placeholder"
    )
    docx_file.write_bytes(
        b"placeholder"
    )
    ignored_file.write_text(
        "ignored",
        encoding="utf-8",
    )

    result = collect_document_files(tmp_path)

    assert result == [
        txt_file,
        pdf_file,
        docx_file,
    ]


def test_collect_text_files_can_include_nested_files(
    tmp_path: Path,
) -> None:
    nested_directory = tmp_path / "nested"
    nested_directory.mkdir()

    nested_file = nested_directory / "nested.txt"
    nested_file.write_text(
        "Nested content",
        encoding="utf-8",
    )

    result = collect_text_files(
        tmp_path,
        recursive=True,
    )

    assert result == [nested_file]


def test_collect_text_files_rejects_unsupported_file(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "unsupported.pdf"
    file_path.write_text(
        "Unsupported",
        encoding="utf-8",
    )

    with pytest.raises(ValueError):
        collect_text_files(file_path)


def test_collect_document_files_rejects_unsupported_file(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "unsupported.xlsx"
    file_path.write_text(
        "Unsupported",
        encoding="utf-8",
    )

    with pytest.raises(ValueError):
        collect_document_files(file_path)


def test_collect_text_files_rejects_missing_path(
    tmp_path: Path,
) -> None:
    missing_path = tmp_path / "missing"

    with pytest.raises(FileNotFoundError):
        collect_text_files(missing_path)


def test_collect_document_files_rejects_missing_path(
    tmp_path: Path,
) -> None:
    missing_path = tmp_path / "missing"

    with pytest.raises(FileNotFoundError):
        collect_document_files(missing_path)