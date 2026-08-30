from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from src.chunker import split_into_chunks
from src.database import (
    get_all_documents,
    get_all_source_files,
)
from src.document_loader import (
    read_text_file,
)
from src.ingestion import (
    ingest_document_file,
    ingest_selected_source,
    ingest_text_files,
)


def test_read_text_file_returns_clean_content(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "test.txt"

    file_path.write_text(
        "  STM32 test content  ",
        encoding="utf-8",
    )

    result = read_text_file(
        file_path
    )

    assert result == "STM32 test content"


def test_read_text_file_rejects_missing_file(
    tmp_path: Path,
) -> None:
    file_path = (
        tmp_path / "missing.txt"
    )

    with pytest.raises(
        FileNotFoundError
    ):
        read_text_file(
            file_path
        )


def test_read_text_file_rejects_non_txt_file(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "test.md"

    file_path.write_text(
        "Markdown content",
        encoding="utf-8",
    )

    with pytest.raises(
        ValueError
    ):
        read_text_file(
            file_path
        )


def test_read_text_file_rejects_empty_file(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "empty.txt"

    file_path.write_text(
        "   ",
        encoding="utf-8",
    )

    with pytest.raises(
        ValueError
    ):
        read_text_file(
            file_path
        )


def test_split_into_chunks_uses_configured_limits() -> None:
    text = " ".join(
        f"kelime{i}"
        for i in range(200)
    )

    chunks = split_into_chunks(
        text=text,
        max_characters=100,
        overlap_characters=20,
    )

    assert len(chunks) > 1

    for chunk in chunks:
        assert chunk.strip()
        assert len(chunk) <= 100


def test_ingestion_stores_generated_chunks(
    tmp_path: Path,
) -> None:
    raw_directory = (
        tmp_path / "raw"
    )

    raw_directory.mkdir()

    file_path = (
        raw_directory / "stm32.txt"
    )

    file_path.write_text(
        "STM32 PWM UART GPIO",
        encoding="utf-8",
    )

    database_path = (
        tmp_path / "test.db"
    )

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.1, 0.2, 0.3],
        ],
    ):
        inserted_count = ingest_text_files(
            reset_database=True,
            raw_data_directory=raw_directory,
            database_path=database_path,
        )

    documents = get_all_documents(
        database_path
    )

    assert inserted_count == 1
    assert len(documents) == 1

    assert (
        documents[0]["source"]
        == "stm32.txt"
    )

    assert (
        documents[0]["content"]
        == "STM32 PWM UART GPIO"
    )

    assert documents[0]["embedding"] == [
        0.1,
        0.2,
        0.3,
    ]


def test_ingestion_rejects_missing_directory(
    tmp_path: Path,
) -> None:
    raw_directory = (
        tmp_path / "missing"
    )

    database_path = (
        tmp_path / "test.db"
    )

    with pytest.raises(
        FileNotFoundError
    ):
        ingest_text_files(
            raw_data_directory=raw_directory,
            database_path=database_path,
        )


def test_ingestion_rejects_directory_without_txt_files(
    tmp_path: Path,
) -> None:
    raw_directory = (
        tmp_path / "raw"
    )

    raw_directory.mkdir()

    (
        raw_directory / "notes.md"
    ).write_text(
        "Markdown",
        encoding="utf-8",
    )

    database_path = (
        tmp_path / "test.db"
    )

    with pytest.raises(
        RuntimeError
    ):
        ingest_text_files(
            raw_data_directory=raw_directory,
            database_path=database_path,
        )


def test_ingestion_creates_source_manifest(
    tmp_path: Path,
) -> None:
    raw_directory = (
        tmp_path / "raw"
    )

    raw_directory.mkdir()

    file_path = (
        raw_directory / "stm32.txt"
    )

    file_path.write_text(
        "STM32 PWM UART GPIO",
        encoding="utf-8",
    )

    database_path = (
        tmp_path / "test.db"
    )

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.1, 0.2, 0.3],
        ],
    ):
        ingest_text_files(
            reset_database=True,
            raw_data_directory=raw_directory,
            database_path=database_path,
        )

    records = get_all_source_files(
        database_path
    )

    assert len(records) == 1

    assert (
        records[0]["source"]
        == "stm32.txt"
    )

    assert records[0]["file_hash"]

    assert isinstance(
        records[0]["modified_at"],
        float,
    )


def test_manifest_hash_changes_when_file_changes(
    tmp_path: Path,
) -> None:
    raw_directory = (
        tmp_path / "raw"
    )

    raw_directory.mkdir()

    file_path = (
        raw_directory / "stm32.txt"
    )

    file_path.write_text(
        "Version 1",
        encoding="utf-8",
    )

    database_path = (
        tmp_path / "test.db"
    )

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.1, 0.2],
        ],
    ):
        ingest_text_files(
            reset_database=True,
            raw_data_directory=raw_directory,
            database_path=database_path,
        )

    first_record = (
        get_all_source_files(
            database_path
        )[0]
    )

    file_path.write_text(
        "Version 2",
        encoding="utf-8",
    )

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.3, 0.4],
        ],
    ):
        ingest_text_files(
            reset_database=True,
            raw_data_directory=raw_directory,
            database_path=database_path,
        )

    second_record = (
        get_all_source_files(
            database_path
        )[0]
    )

    assert (
        first_record["file_hash"]
        != second_record["file_hash"]
    )


def test_ingest_document_file_supports_txt(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "knowledge.txt"

    file_path.write_text(
        "LocalMind supports private document retrieval.",
        encoding="utf-8",
    )

    database_path = tmp_path / "test.db"

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.1, 0.2, 0.3],
        ],
    ):
        inserted_count = ingest_document_file(
            file_path=file_path,
            database_path=database_path,
        )

    documents = get_all_documents(
        database_path
    )

    assert inserted_count == 1
    assert len(documents) == 1

    assert (
        documents[0]["source"]
        == "knowledge.txt"
    )

    assert (
        "LocalMind supports private document retrieval."
        in documents[0]["content"]
    )


def test_ingest_document_file_supports_docx(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "knowledge.docx"

    document = Document()
    document.add_paragraph(
        "DOCX-482 is a LocalMind test identifier."
    )
    document.save(file_path)

    database_path = tmp_path / "test.db"

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.2, 0.3, 0.4],
        ],
    ):
        inserted_count = ingest_document_file(
            file_path=file_path,
            database_path=database_path,
        )

    documents = get_all_documents(
        database_path
    )

    assert inserted_count == 1
    assert len(documents) == 1

    assert (
        documents[0]["source"]
        == "knowledge.docx"
    )

    assert (
        "DOCX-482"
        in documents[0]["content"]
    )


def test_ingest_document_file_supports_pdf_extracted_text(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "knowledge.pdf"

    file_path.write_bytes(
        b"placeholder-pdf"
    )

    database_path = tmp_path / "test.db"

    with (
        patch(
            "src.ingestion.read_document_file",
            return_value=(
                "PDF-731 is a LocalMind PDF test identifier."
            ),
        ),
        patch(
            "src.ingestion.generate_embeddings",
            return_value=[
                [0.3, 0.4, 0.5],
            ],
        ),
    ):
        inserted_count = ingest_document_file(
            file_path=file_path,
            database_path=database_path,
        )

    documents = get_all_documents(
        database_path
    )

    assert inserted_count == 1
    assert len(documents) == 1

    assert (
        documents[0]["source"]
        == "knowledge.pdf"
    )

    assert (
        "PDF-731"
        in documents[0]["content"]
    )


def test_ingest_document_file_rejects_unsupported_type(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "knowledge.xlsx"

    file_path.write_text(
        "Unsupported document",
        encoding="utf-8",
    )

    database_path = tmp_path / "test.db"

    with pytest.raises(ValueError):
        ingest_document_file(
            file_path=file_path,
            database_path=database_path,
        )


def test_selected_source_uses_external_txt_source_name(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "notes.txt"

    file_path.write_text(
        "External TXT document",
        encoding="utf-8",
    )

    database_path = tmp_path / "test.db"

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.1, 0.2],
        ],
    ):
        result = ingest_selected_source(
            source_path=file_path,
            database_path=database_path,
        )

    assert result["file_count"] == 1

    assert result["sources"] == [
        "external/notes.txt"
    ]

    documents = get_all_documents(
        database_path
    )

    assert (
        documents[0]["source"]
        == "external/notes.txt"
    )


def test_selected_source_uses_external_docx_source_name(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "lecture.docx"

    document = Document()
    document.add_paragraph(
        "DOCX external source test."
    )
    document.save(file_path)

    database_path = tmp_path / "test.db"

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.2, 0.4],
        ],
    ):
        result = ingest_selected_source(
            source_path=file_path,
            database_path=database_path,
        )

    assert result["sources"] == [
        "external/lecture.docx"
    ]

    documents = get_all_documents(
        database_path
    )

    assert (
        documents[0]["source"]
        == "external/lecture.docx"
    )


def test_selected_source_uses_external_pdf_source_name(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "lecture.pdf"

    file_path.write_bytes(
        b"placeholder-pdf"
    )

    database_path = tmp_path / "test.db"

    with (
        patch(
            "src.ingestion.read_document_file",
            return_value=(
                "PDF external source test."
            ),
        ),
        patch(
            "src.ingestion.generate_embeddings",
            return_value=[
                [0.3, 0.5],
            ],
        ),
    ):
        result = ingest_selected_source(
            source_path=file_path,
            database_path=database_path,
        )

    assert result["sources"] == [
        "external/lecture.pdf"
    ]

    documents = get_all_documents(
        database_path
    )

    assert (
        documents[0]["source"]
        == "external/lecture.pdf"
    )


def test_selected_source_accepts_custom_docx_name(
    tmp_path: Path,
) -> None:
    temporary_file = tmp_path / "temporary.docx"

    document = Document()
    document.add_paragraph(
        "Original DOCX upload test."
    )
    document.save(temporary_file)

    database_path = tmp_path / "test.db"

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.4, 0.6],
        ],
    ):
        result = ingest_selected_source(
            source_path=temporary_file,
            database_path=database_path,
            external_source_name=(
                "project_report.docx"
            ),
        )

    assert result["sources"] == [
        "external/project_report.docx"
    ]


def test_selected_source_accepts_custom_pdf_name(
    tmp_path: Path,
) -> None:
    temporary_file = tmp_path / "temporary.pdf"

    temporary_file.write_bytes(
        b"placeholder-pdf"
    )

    database_path = tmp_path / "test.db"

    with (
        patch(
            "src.ingestion.read_document_file",
            return_value=(
                "Uploaded PDF document."
            ),
        ),
        patch(
            "src.ingestion.generate_embeddings",
            return_value=[
                [0.5, 0.7],
            ],
        ),
    ):
        result = ingest_selected_source(
            source_path=temporary_file,
            database_path=database_path,
            external_source_name=(
                "project_report.pdf"
            ),
        )

    assert result["sources"] == [
        "external/project_report.pdf"
    ]


def test_selected_source_rejects_unsupported_custom_name(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "temporary.txt"

    file_path.write_text(
        "Temporary upload",
        encoding="utf-8",
    )

    database_path = tmp_path / "test.db"

    with pytest.raises(ValueError):
        ingest_selected_source(
            source_path=file_path,
            database_path=database_path,
            external_source_name=(
                "malicious.exe"
            ),
        )


def test_selected_source_replaces_existing_source(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "replace.txt"

    database_path = tmp_path / "test.db"

    file_path.write_text(
        "Version one",
        encoding="utf-8",
    )

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.1, 0.2],
        ],
    ):
        first_result = ingest_selected_source(
            source_path=file_path,
            database_path=database_path,
        )

    assert (
        first_result["inserted_chunks"]
        == 1
    )

    file_path.write_text(
        "Version two",
        encoding="utf-8",
    )

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[
            [0.3, 0.4],
        ],
    ):
        second_result = ingest_selected_source(
            source_path=file_path,
            database_path=database_path,
        )

    documents = get_all_documents(
        database_path
    )

    assert (
        second_result["inserted_chunks"]
        == 1
    )

    assert len(documents) == 1

    assert (
        documents[0]["content"]
        == "Version two"
    )

    assert (
        documents[0]["source"]
        == "external/replace.txt"
    )


def test_ingest_document_file_rejects_embedding_count_mismatch(
    tmp_path: Path,
) -> None:
    file_path = tmp_path / "document.txt"

    file_path.write_text(
        "Embedding count validation.",
        encoding="utf-8",
    )

    database_path = tmp_path / "test.db"

    with patch(
        "src.ingestion.generate_embeddings",
        return_value=[],
    ):
        with pytest.raises(
            RuntimeError,
            match=(
                "Chunk and embedding counts "
                "do not match"
            ),
        ):
            ingest_document_file(
                file_path=file_path,
                database_path=database_path,
            )